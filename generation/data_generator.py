import json
import os
import sqlite3
import random
from datetime import datetime, timedelta
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# ── third-party ──────────────────────────────────────────────────────────────
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
import openpyxl
from faker import Faker

# ── config ───────────────────────────────────────────────────────────────────
BASE_DIR   = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BIBLE_PATH = os.path.join(BASE_DIR, "world_bible.json")
RAW        = os.path.join(BASE_DIR, "data", "raw")

fake = Faker("en_IN")
random.seed(42)

# ── load world bible ─────────────────────────────────────────────────────────
with open(BIBLE_PATH, encoding="utf-8") as f:
    WB = json.load(f)

CO   = WB["company"]
DEPT = {d["id"]: d for d in WB["departments"]}
PPL  = WB["people"]
PROJ = WB["projects"]
VEN  = WB["vendors"]
POL  = WB["policies"]
AUD  = WB["audit_records"]
INV  = WB["invoices"]
PC   = WB["pressure_cooker_products"]
CW   = WB["cookware_products"]
KE   = WB["kitchen_electricals"]
ACC  = WB["accessories"]
INTERNS = WB["interns"]
PROD_LOG = WB["production_log"]
SVC  = WB["service_centres"]
EXP  = WB["export_markets"]
CB   = WB["cookbooks_manuals"]
GP   = WB["gift_packs"]


# ─────────────────────────────────────────────────────────────────────────────
# HELPER UTILITIES
# ─────────────────────────────────────────────────────────────────────────────

def person(id_):
    return next(p for p in PPL if p["id"] == id_)

def vendor(id_):
    return next(v for v in VEN if v["id"] == id_)

def project(id_):
    return next(p for p in PROJ if p["id"] == id_)

def rand_date(start="2024-01-01", end="2025-12-31"):
    s = datetime.strptime(start, "%Y-%m-%d")
    e = datetime.strptime(end,   "%Y-%m-%d")
    return (s + timedelta(days=random.randint(0, (e - s).days))).strftime("%Y-%m-%d")


# ─────────────────────────────────────────────────────────────────────────────
# PDF GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def safe_text(s, max_len=300):
    """Clean text so fpdf2 never chokes on it: strip weird unicode, collapse long unbroken words."""
    if not s:
        return ""
    s = str(s).strip()
    # Replace problem characters fpdf2's core font can't render
    s = (s.replace("\u2018", "'").replace("\u2019", "'")
           .replace("\u201c", '"').replace("\u201d", '"')
           .replace("\u2013", "-").replace("\u2014", "-")
           .replace("\u2026", "..."))
    s = s.encode("latin-1", "ignore").decode("latin-1")
    # Break up any single "word" longer than 40 chars so multi_cell can wrap it
    words = s.split(" ")
    fixed_words = []
    for w in words:
        if len(w) > 40:
            w = " ".join(w[i:i+40] for i in range(0, len(w), 40))
        fixed_words.append(w)
    s = " ".join(fixed_words)
    return s[:max_len] if max_len else s


class HawkinsPDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(180, 30, 30)
        self.cell(0, 8, safe_text(CO["name"], 100), new_x="LMARGIN", new_y="NEXT")
        self.set_text_color(0, 0, 0)
        self.set_font("Helvetica", "", 8)
        self.cell(0, 5, safe_text(CO["registered_office"], 150), new_x="LMARGIN", new_y="NEXT")
        self.ln(2)
        self.set_draw_color(180, 30, 30)
        self.set_line_width(0.5)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(4)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(120, 120, 120)
        self.cell(0, 10, safe_text(f"Page {self.page_no()} | {CO['name']} - Confidential", 150), align="C")

    def title_block(self, title, subtitle="", date=""):
        self.set_font("Helvetica", "B", 14)
        self.set_text_color(30, 30, 30)
        self.set_x(self.l_margin)
        self.multi_cell(0, 8, safe_text(title, 200), align="L")
        if subtitle:
            self.set_font("Helvetica", "I", 11)
            self.set_text_color(80, 80, 80)
            self.set_x(self.l_margin)
            self.multi_cell(0, 6, safe_text(subtitle, 200), align="L")
        if date:
            self.set_font("Helvetica", "", 9)
            self.set_text_color(100, 100, 100)
            self.set_x(self.l_margin)
            self.cell(0, 6, safe_text(f"Date: {date}", 100), new_x="LMARGIN", new_y="NEXT", align="L")
        self.ln(4)

    def section(self, heading):
        self.set_font("Helvetica", "B", 11)
        self.set_text_color(180, 30, 30)
        self.set_x(self.l_margin)
        self.cell(0, 7, safe_text(heading, 150), new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(180, 30, 30)
        self.set_line_width(0.3)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(3)
        self.set_text_color(0, 0, 0)

    def body(self, text):
        self.set_font("Helvetica", "", 10)
        self.set_text_color(30, 30, 30)
        self.set_x(self.l_margin)
        self.multi_cell(0, 6, safe_text(text, 3000))
        self.ln(2)

    def kv(self, key, value):
        self.set_font("Helvetica", "B", 10)
        self.set_x(self.l_margin)
        self.cell(55, 6, safe_text(str(key), 60) + ":", new_x="RIGHT", new_y="TOP")
        self.set_font("Helvetica", "", 10)
        self.multi_cell(0, 6, safe_text(str(value), 800))

    def table_row(self, cols, widths, bold=False):
        self.set_font("Helvetica", "B" if bold else "", 9)
        self.set_x(self.l_margin)
        for col, w in zip(cols, widths):
            self.cell(w, 6, safe_text(str(col), 30), border=1)
        self.ln()


def save_pdf(pdf, filename):
    path = os.path.join(RAW, "pdfs", filename)
    pdf.output(path)
    print(f"  [PDF] {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# PDF DOCUMENTS
# ─────────────────────────────────────────────────────────────────────────────

def pdf_company_overview():
    pdf = HawkinsPDF(); pdf.add_page()
    pdf.title_block("Company Overview", CO["name"], "June 2025")
    pdf.section("About Hawkins Cookers Limited")
    pdf.body(
        f"{CO['name']} was founded in {CO['founded']} by {CO['founder']} in technical "
        f"collaboration with L.G. Hawkins of England, with a starting capital of Rs 20,000. "
        f"Today the company employs approximately {CO['employees']} people and operates "
        f"{len(CO['factories'])} manufacturing plants and {len(CO['warehouses'])} warehouses across India."
    )
    pdf.body(
        f"Hawkins is the market leader in pressure cookers in India and has exported products "
        f"since {CO['export_since']} to countries across all six continents. "
        f"Worldwide the company has sold over {CO['total_units_sold_worldwide']} units of pressure cookers and cookware."
    )
    pdf.section("Manufacturing Locations")
    for f in CO["factories"]:
        pdf.kv(f["location"], f"{f['type']} | Est. {f['established']} | Capacity: {f['capacity_units_per_day']} units/day")
    pdf.section("Certifications")
    for c in CO["certifications"]:
        pdf.body(f"• {c}")
    pdf.section("Mission Statement")
    pdf.body(CO["mission"])
    save_pdf(pdf, "01_Company_Overview.pdf")


def pdf_project_report(proj):
    pdf = HawkinsPDF(); pdf.add_page()
    lead = next(p for p in PPL if p["name"] == proj["lead"])
    pdf.title_block(f"{proj['name']} — Project Status Report",
                    DEPT[proj['dept']]['name'], rand_date())
    pdf.section("Project Summary")
    pdf.kv("Project ID",    proj["id"])
    pdf.kv("Project Name",  proj["name"])
    pdf.kv("Department",    DEPT[proj['dept']]['name'])
    pdf.kv("Lead",          proj["lead"])
    pdf.kv("Status",        proj["status"])
    pdf.kv("Budget",        f"Rs {proj['budget']:,}")
    pdf.kv("Spent to Date", f"Rs {proj['spent']:,}")
    pdf.kv("Start Date",    proj["start"])
    pdf.kv("End Date",      proj["end"])
    pdf.section("Project Description")
    pdf.body(proj["description"])
    pdf.section("Phase-wise Progress")
    pdf.table_row(["Phase", "Name", "Status", "Target End"], [15, 75, 30, 40], bold=True)
    for ph in proj["phases"]:
        pdf.table_row([ph["phase"], ph["name"], ph["status"], ph["end"]], [15, 75, 30, 40])
    pdf.section("Budget Utilisation")
    utilisation = round(proj["spent"] / proj["budget"] * 100, 1)
    pdf.body(f"Total approved budget: Rs {proj['budget']:,}\n"
             f"Spent to date: Rs {proj['spent']:,}\n"
             f"Remaining: Rs {proj['budget'] - proj['spent']:,}\n"
             f"Utilisation: {utilisation}%")
    pdf.section("Lead Contact")
    pdf.kv("Name",  lead["name"])
    pdf.kv("Role",  lead["role"])
    pdf.kv("Email", lead["email"])
    pdf.kv("Phone", lead["phone"])
    fname = f"02_Project_{proj['id']}_{proj['name'].replace(' ', '_')}_Report.pdf"
    save_pdf(pdf, fname)


def pdf_vendor_evaluation(ven):
    pdf = HawkinsPDF(); pdf.add_page()
    pdf.title_block(f"Vendor Evaluation Report — {ven['name']}",
                    "Procurement Department", rand_date())
    pdf.section("Vendor Details")
    pdf.kv("Vendor ID",       ven["id"])
    pdf.kv("Company Name",    ven["name"])
    pdf.kv("Material",        ven["material"])
    pdf.kv("Location",        ven["location"])
    pdf.kv("Contact Person",  ven["contact"])
    pdf.kv("Email",           ven["email"])
    pdf.kv("Contract Value",  f"Rs {ven['contract_value']:,}")
    pdf.kv("Performance Score", f"{ven['rating']}/100")
    pdf.kv("Vendor Since",    str(ven["since"]))
    pdf.section("Evaluation Criteria")
    pdf.table_row(["Criteria", "Score (out of 25)", "Remarks"], [70, 45, 65], bold=True)
    criteria = [
        ("Quality of materials supplied",    min(ven["rating"] // 4, 25), "Consistent with spec"),
        ("On-time delivery performance",     min((ven["rating"] - 2) // 4, 25), "Meets schedule"),
        ("Pricing competitiveness",          min((ven["rating"] - 1) // 4, 25), "Within budget"),
        ("Communication and responsiveness", min(ven["rating"] // 4, 25), "Satisfactory"),
    ]
    for c, s, r in criteria:
        pdf.table_row([c, s, r], [70, 45, 65])
    pdf.section("Recommendation")
    rec = "APPROVED — Recommended for contract renewal" if ven["rating"] >= 85 \
          else "CONDITIONAL — Improvement plan required"
    pdf.body(rec)
    pdf.body(f"Evaluated by: {person('E003')['name']}, {person('E003')['role']}")
    fname = f"03_Vendor_Evaluation_{ven['id']}_{ven['name'].replace(' ', '_')[:20]}.pdf"
    save_pdf(pdf, fname)


def pdf_qc_inspection_report():
    pdf = HawkinsPDF(); pdf.add_page()
    pdf.title_block("QC Inspection Report — Thane Plant",
                    "Quality Control Department", rand_date())
    pdf.section("Inspection Overview")
    pdf.body(
        "This report covers the quality control inspection conducted at the Thane manufacturing "
        "plant for the AlphaCook pressure cooker production lines. Inspections were carried out "
        "as per the QC Inspection Standards Manual (POL003, v4.0) effective March 2024."
    )
    pdf.section("Defect Summary by Product Line")
    pdf.table_row(["Product Line", "Batch", "Units Inspected", "Defects Found", "Defect Rate %"], [40,30,35,30,30], bold=True)
    rows = [
        ("Hawkins Classic 3L", "CL-2025-041", 500, 4, "0.80%"),
        ("Hawkins Contura 3L", "HC-2025-012", 300, 6, "2.00%"),
        ("Hawkins Instaa 2L",  "II-2025-008", 400, 3, "0.75%"),
        ("Futura HA 4L",       "FH-2025-019", 250, 2, "0.80%"),
        ("Futura SS 5L",       "FS-2025-003", 200, 5, "2.50%"),
    ]
    for r in rows:
        pdf.table_row(list(r), [40,30,35,30,30])
    pdf.section("Critical Findings")
    pdf.body(
        "1. Batch HC-2025-012 (Hawkins Contura 3L): Ceramic nonstick coating adhesion found "
        "below specification on 6 units. Root cause: temperature variance on coating line 2. "
        "Corrective action: Recalibration of oven temperature sensors. Status: Resolved.\n\n"
        "2. Batch FS-2025-003 (Futura SS 5L): Lid sealing ring fitment loose on 5 units. "
        "Root cause: Dimensional variance in gasket batch from SafeSeal Gaskets. "
        "Corrective action: Vendor notified, replacement batch ordered. Status: Pending."
    )
    pdf.section("Sign-off")
    pdf.kv("Inspected by", f"{person('E002')['name']}, {person('E002')['role']}")
    pdf.kv("Date",         rand_date())
    save_pdf(pdf, "04_QC_Inspection_Report_Thane_Plant.pdf")


def pdf_hr_policy(pol):
    pdf = HawkinsPDF(); pdf.add_page()
    pdf.title_block(pol["title"], f"Department: {DEPT[pol['dept']]['name']}",
                    f"Effective: {pol['effective']} | {pol['revision']}")
    pdf.section("Policy Overview")
    pdf.body(pol["desc"])
    if pol["id"] == "POL001":
        pdf.section("Leave Entitlement for Interns")
        pdf.body(
            "All interns at Hawkins Cookers Limited are entitled to the following leave during "
            "their internship period:\n\n"
            "• Casual Leave: 1 day per month (non-cumulative)\n"
            "• Medical Leave: Up to 2 days per internship period, with medical certificate\n"
            "• Public Holidays: As per Maharashtra Government calendar\n\n"
            "Leave must be applied at least 48 hours in advance through the HR portal, "
            "except in cases of medical emergency. Unapproved absence will be treated as "
            "Loss of Pay (LOP) and will be deducted from the internship stipend."
        )
        pdf.section("Working Hours")
        pdf.body(
            "Standard working hours: 9:00 AM to 6:00 PM, Monday to Friday.\n"
            "Lunch break: 1:00 PM to 2:00 PM.\n"
            "Interns are expected to maintain 90% attendance throughout the internship period.\n"
            "Work from home is permitted on a case-by-case basis with mentor approval."
        )
        pdf.section("Code of Conduct")
        pdf.body(
            "All interns must:\n"
            "• Maintain confidentiality of all company data and processes\n"
            "• Sign the Non-Disclosure Agreement (NDA) on Day 1\n"
            "• Not share any proprietary information on social media\n"
            "• Dress formally (business casual) on all working days\n"
            "• Carry their intern ID card at all times within company premises"
        )
        pdf.section("Stipend and Benefits")
        pdf.body(
            "Monthly stipend: Rs 15,000 (subject to 90% attendance)\n"
            "Completion bonus: Rs 5,000 on successful project submission\n"
            "Certificate: Issued within 15 working days of completion\n"
            "Lunch: Subsidised canteen access at Rs 50 per meal"
        )
    pdf.section("Issued By")
    hr = person("E004")
    pdf.kv("HR Manager", hr["name"])
    pdf.kv("Email",      hr["email"])
    fname = f"05_Policy_{pol['id']}_{pol['title'].replace(' ', '_')[:30]}.pdf"
    save_pdf(pdf, fname)


def pdf_annual_production_review():
    pdf = HawkinsPDF(); pdf.add_page()
    pdf.title_block("Annual Production Review 2025", "Operations Department", "January 2026")
    pdf.section("Executive Summary")
    total = sum(r["units_produced"] for r in PROD_LOG)
    rejected = sum(r["units_rejected"] for r in PROD_LOG)
    pdf.body(
        f"Total units produced across all plants in 2025: {total:,}\n"
        f"Total units rejected: {rejected:,}\n"
        f"Overall defect rate: {round(rejected/total*100,2)}%\n"
        f"Target defect rate: <1.0%\n"
        f"Status: {'MET' if rejected/total*100 < 1.0 else 'MISSED — improvement plan required'}"
    )
    pdf.section("Monthly Production Log — All Plants")
    pdf.table_row(["Month", "Plant", "Product Line", "Produced", "Rejected", "Defect %"],
                  [25, 20, 45, 25, 25, 25], bold=True)
    for r in PROD_LOG:
        pdf.table_row([r["month"], r["plant"], r["product_line"],
                       r["units_produced"], r["units_rejected"], r["defect_rate_pct"]],
                      [25, 20, 45, 25, 25, 25])
    pdf.section("Plant-wise Summary")
    plants = {}
    for r in PROD_LOG:
        p = r["plant"]
        plants.setdefault(p, {"produced": 0, "rejected": 0})
        plants[p]["produced"]  += r["units_produced"]
        plants[p]["rejected"]  += r["units_rejected"]
    for p, d in plants.items():
        pdf.body(f"{p}: {d['produced']:,} units produced, {d['rejected']} rejected "
                 f"({round(d['rejected']/d['produced']*100,2)}% defect rate)")
    save_pdf(pdf, "06_Annual_Production_Review_2025.pdf")


def pdf_export_report():
    pdf = HawkinsPDF(); pdf.add_page()
    pdf.title_block("Export Operations Report 2025",
                    "Export & International Department", rand_date())
    pdf.section("Overview")
    pdf.body(
        f"{CO['name']} has been exporting products since {CO['export_since']} and currently "
        f"serves markets across all six continents. All exported pressure cookers carry UL listing "
        f"(USA) and CE certification (Europe). From February 2021, all products also carry the "
        f"BIS/ISI mark as mandated by the Government of India."
    )
    pdf.section("Market-wise Performance")
    pdf.table_row(["Region", "Key Countries", "Key Products"], [35, 75, 70], bold=True)
    for m in EXP:
        pdf.table_row([m["region"], ", ".join(m["countries"]), m["key_product"]],
                      [35, 75, 70])
    pdf.section("Export Manager")
    exp = person("E010")
    pdf.kv("Name",  exp["name"])
    pdf.kv("Role",  exp["role"])
    pdf.kv("Email", exp["email"])
    save_pdf(pdf, "07_Export_Operations_Report_2025.pdf")


def pdf_product_spec(product, category="PC"):
    pdf = HawkinsPDF(); pdf.add_page()
    name = product.get("brand", "") + " " + product.get("type", product.get("model", ""))
    pdf.title_block(f"Product Specification Sheet\n{name}",
                    "Research & Development Department", rand_date())
    pdf.section("Product Details")
    for k, v in product.items():
        if k not in ("id",):
            pdf.kv(k.replace("_", " ").title(), str(v))
    pdf.section("Quality Standards")
    pdf.body(
        "All Hawkins products are manufactured to the following standards:\n"
        "• BIS/ISI certification (mandatory from February 1, 2021)\n"
        "• Underwriters Laboratories (UL) listing — USA\n"
        "• CE certification for European markets\n"
        "• Internal QC inspection pass rate: >99%\n"
        "• 5-year guarantee on pressure cookers"
    )
    pdf.section("R&D Sign-off")
    rd = person("E001")
    pdf.kv("Approved by", f"{rd['name']}, {rd['role']}")
    pdf.kv("Email",       rd["email"])
    idx = product["id"]
    save_pdf(pdf, f"08_Product_Spec_{idx}_{name.replace(' ', '_')[:25]}.pdf")


def pdf_it_project_report():
    proj = project("P004")
    intern = next(i for i in INTERNS if i["id"] == "I001")
    pdf = HawkinsPDF(); pdf.add_page()
    pdf.title_block("IT & AI Project Report — Project Horizon",
                    "Information Technology Department", rand_date("2025-06-01","2026-06-30"))
    pdf.section("Project Overview")
    pdf.body(proj["description"])
    pdf.section("AI Knowledge Assistant Sub-Project")
    pdf.body(
        f"As part of Phase 2 of Project Horizon, an intern-driven initiative has been launched "
        f"to build a Local Enterprise Knowledge Assistant. The system is being developed by "
        f"{intern['name']} from {intern['college']} under the mentorship of "
        f"{person('E008')['name']}, IT Head.\n\n"
        f"The system will ingest documents from multiple sources — PDFs, Word documents, "
        f"Excel sheets, emails, and SQL databases — and allow employees to query the knowledge "
        f"base using natural language. All processing happens on-premise with no data leaving "
        f"the company network."
    )
    pdf.section("Technology Stack")
    pdf.body(
        "• Connector modules: pdfplumber, python-docx, openpyxl, sqlite3\n"
        "• Metadata tagging: LLaMA 3.1:8b via Ollama (local inference)\n"
        "• Embeddings: BGE-M3 (sentence-transformers)\n"
        "• Vector database: ChromaDB (persistent, local)\n"
        "• UI: Streamlit dashboard\n"
        "• Hardware: NVIDIA RTX 3050 6GB (local GPU inference)"
    )
    pdf.section("Timeline")
    for ph in proj["phases"]:
        pdf.kv(f"Phase {ph['phase']}: {ph['name']}", f"{ph['status']} — Due {ph['end']}")
    pdf.section("Budget")
    pdf.kv("Approved Budget", f"Rs {proj['budget']:,}")
    pdf.kv("Spent to Date",   f"Rs {proj['spent']:,}")
    save_pdf(pdf, "09_IT_Project_Horizon_Report.pdf")


def pdf_service_centre_report():
    pdf = HawkinsPDF(); pdf.add_page()
    pdf.title_block("Authorised Service Centre Network Report",
                    "Customer Service Department", rand_date())
    pdf.section("Overview")
    pdf.body(
        f"{CO['name']} maintains a network of {len(SVC)} authorised service centres across "
        f"major cities in India. Each centre is equipped to handle warranty claims, repair kits, "
        f"and replacement parts for all Hawkins and Futura products."
    )
    pdf.section("Service Centre Directory")
    pdf.table_row(["ID", "City", "Phone"], [20, 40, 60], bold=True)
    for s in SVC:
        pdf.table_row([s["id"], s["city"], s["phone"]], [20, 40, 60])
    pdf.section("Warranty Policy")
    pdf.body(
        "All Hawkins pressure cookers carry a 5-year guarantee from date of purchase. "
        "Cookware is guaranteed as specified on packaging. Claims must be made at any "
        "authorised service centre with original purchase receipt. The Hawkins Repair Kit "
        "allows customers to self-repair cookers without visiting a service centre."
    )
    pdf.section("Customer Service Head")
    cs = person("E009")
    pdf.kv("Name",  cs["name"])
    pdf.kv("Email", cs["email"])
    save_pdf(pdf, "10_Service_Centre_Network_Report.pdf")


def pdf_finance_budget():
    pdf = HawkinsPDF(); pdf.add_page()
    pdf.title_block("Finance — Q1 2025 Budget Summary",
                    "Finance Department", "April 2025")
    pdf.section("Budget vs Actuals — Q1 2025")
    pdf.table_row(["Department", "Budget (Rs)", "Spent (Rs)", "Variance"], [50,40,40,40], bold=True)
    rows = [
        ("R&D",            "15,00,000", "11,20,000", "+3,80,000"),
        ("Quality Control", "8,00,000",  "7,50,000",  "+50,000"),
        ("Procurement",    "12,00,000", "9,80,000",  "+2,20,000"),
        ("IT (Horizon)",    "6,00,000",  "4,20,000",  "+1,80,000"),
        ("HR",              "4,00,000",  "3,85,000",  "+15,000"),
        ("Operations",     "45,00,000", "44,20,000", "+80,000"),
        ("Sales & Mktg",   "18,00,000", "16,50,000", "+1,50,000"),
    ]
    for r in rows:
        pdf.table_row(list(r), [50,40,40,40])
    pdf.section("Key Observations")
    pdf.body(
        "1. All departments are within approved budget for Q1 2025.\n"
        "2. R&D underspend due to Aurora Phase 3 delay in vendor tooling.\n"
        "3. Operations near full utilisation — budget revision may be needed in Q2.\n"
        "4. IT budget includes Rs 4,20,000 for ERP assessment and AI prototype work."
    )
    pdf.section("Approved by")
    fin = person("E005")
    pdf.kv("Finance Director", fin["name"])
    pdf.kv("Email",            fin["email"])
    save_pdf(pdf, "11_Finance_Q1_2025_Budget_Summary.pdf")


def pdf_intern_onboarding():
    pdf = HawkinsPDF(); pdf.add_page()
    pdf.title_block("Intern Onboarding Guide — Summer 2026",
                    "Human Resources Department", "June 2026")
    pdf.section("Welcome")
    pdf.body(
        f"Welcome to {CO['name']}! We are delighted to have you join us as an intern. "
        f"This guide will help you navigate your first week and make the most of your "
        f"internship experience."
    )
    pdf.section("Intern Cohort — Summer 2026")
    pdf.table_row(["Name", "College", "Department", "Project"], [40,40,35,55], bold=True)
    for i in INTERNS:
        dept_name = DEPT[i["dept"]]["name"][:20]
        pdf.table_row([i["name"], i["college"], dept_name, i["project"][:25]], [40,40,35,55])
    pdf.section("Day 1 Checklist")
    pdf.body(
        "• Collect your ID card from Security (Gate 1, Ground Floor)\n"
        "• Sign NDA and joining documents with HR (HR Office, 3rd Floor)\n"
        "• Set up your laptop with IT (IT Help Desk, 4th Floor)\n"
        "• Meet your mentor and collect your project brief\n"
        "• Attend orientation session at 11:00 AM, Conference Room B"
    )
    pdf.section("Important Contacts")
    for e in [person("E004"), person("E008"), person("E007")]:
        pdf.kv(e["role"], f"{e['name']} | {e['email']}")
    save_pdf(pdf, "12_Intern_Onboarding_Guide_2026.pdf")


def pdf_safety_compliance():
    pdf = HawkinsPDF(); pdf.add_page()
    pdf.title_block("Safety Compliance Certification Report 2025",
                    "Quality Control & Operations", rand_date())
    pdf.section("Certification Status")
    pdf.table_row(["Certification", "Body", "Status", "Valid Until"], [50,50,30,40], bold=True)
    certs = [
        ("BIS/ISI Mark",              "Bureau of Indian Standards", "Active", "Mar 2027"),
        ("UL Listing",                "Underwriters Laboratories USA", "Active", "Dec 2026"),
        ("CE Certification",          "EU Notified Body", "Active", "Jun 2027"),
        ("ISO 9001:2015",             "Bureau Veritas", "Active", "Sep 2026"),
        ("Factory License — Thane",   "Maharashtra Govt", "Active", "Dec 2025"),
        ("Factory License — Hoshiarpur","Punjab Govt",   "Active", "Mar 2026"),
        ("Factory License — Jaunpur", "UP Govt",         "Active", "Jun 2026"),
    ]
    for c in certs:
        pdf.table_row(list(c), [50,50,30,40])
    pdf.section("Compliance Notes")
    pdf.body(
        "All Hawkins pressure cookers manufactured and sold in India carry the mandatory "
        "BIS/ISI stamp as required by the Government of India from February 1, 2021. "
        "Exported products to USA carry UL listing and European exports carry CE mark. "
        "Annual re-certification audits are conducted by approved third-party agencies."
    )
    save_pdf(pdf, "13_Safety_Compliance_Certification_2025.pdf")


def pdf_product_catalog_summary():
    pdf = HawkinsPDF(); pdf.add_page()
    pdf.title_block("Product Range Summary — Pressure Cookers 2025",
                    "Sales & Marketing Department", "January 2025")
    pdf.section("Pressure Cooker Families")
    families = {}
    for p in PC:
        b = p["brand"]
        families.setdefault(b, []).append(p)
    for brand, products in families.items():
        pdf.section(brand)
        models = ", ".join(f"{p['model']} (Rs {p['mrp']})" for p in products[:5])
        pdf.body(f"Models: {models}")
        pdf.body(products[0]["desc"])
    save_pdf(pdf, "14_Product_Catalog_Summary_PC_2025.pdf")


def pdf_vendor_contract(ven):
    pdf = HawkinsPDF(); pdf.add_page()
    pdf.title_block(f"Supply Agreement — {ven['name']}",
                    "Procurement Department", rand_date())
    pdf.section("Agreement Details")
    pdf.kv("Vendor",           ven["name"])
    pdf.kv("Material",         ven["material"])
    pdf.kv("Contract Value",   f"Rs {ven['contract_value']:,} per annum")
    pdf.kv("Contract Period",  "January 2025 to December 2025")
    pdf.kv("Payment Terms",    "Net 30 days from invoice date")
    pdf.kv("Delivery Terms",   "DDP — Delivered Duty Paid to Hawkins plant")
    pdf.section("Quality Requirements")
    pdf.body(
        "All materials supplied must conform to Hawkins QC Inspection Standards (POL003). "
        "Rejection rate must not exceed 0.5%. In case of batch rejection, vendor must "
        "replace within 7 working days at no additional cost."
    )
    pdf.section("Authorised Signatories")
    proc = person("E003")
    fin  = person("E005")
    pdf.kv("Hawkins — Procurement", f"{proc['name']}, {proc['role']}")
    pdf.kv("Hawkins — Finance",     f"{fin['name']}, {fin['role']}")
    pdf.kv("Vendor Contact",        f"{ven['contact']}, {ven['email']}")
    fname = f"15_Supply_Agreement_{ven['id']}_{ven['name'].replace(' ','_')[:20]}.pdf"
    save_pdf(pdf, fname)


def pdf_audit_summary():
    pdf = HawkinsPDF(); pdf.add_page()
    pdf.title_block("Audit Records Summary 2024-2025",
                    "Quality Control Department", "January 2026")
    pdf.section("All Audit Records")
    pdf.table_row(["Audit ID", "Project", "Date", "Status"], [25,30,30,40], bold=True)
    for a in AUD:
        pdf.table_row([a["audit_id"], a["project_id"], a["date"], a["status"]], [25,30,30,40])
    pdf.section("Detailed Findings")
    for a in AUD:
        pdf.section(f"{a['audit_id']} — {a['date']}")
        pdf.kv("Project",   a["project_id"])
        pdf.kv("Plant",     a["plant"])
        pdf.kv("Auditor",   a["auditor"])
        pdf.kv("Finding",   a["finding"])
        pdf.kv("Status",    a["status"])
        if a["resolution_date"]:
            pdf.kv("Resolved", a["resolution_date"])
    save_pdf(pdf, "16_Audit_Records_Summary_2024_2025.pdf")


# ─────────────────────────────────────────────────────────────────────────────
# WORD DOCUMENT GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = RGBColor(180, 30, 30)


def add_para(doc, text):
    doc.add_paragraph(text)


def save_docx(doc, filename):
    path = os.path.join(RAW, "docx", filename)
    doc.save(path)
    print(f"  [DOCX] {filename}")


def docx_executive_summary(proj):
    doc = Document()
    doc.add_heading(CO["name"], 0)
    add_heading(doc, f"Executive Summary — {proj['name']}")
    add_para(doc, f"Project: {proj['name']} | Department: {DEPT[proj['dept']]['name']} | Lead: {proj['lead']}")
    add_heading(doc, "Project Overview", 2)
    add_para(doc, proj["description"])
    add_heading(doc, "Current Status", 2)
    add_para(doc, f"Status: {proj['status']}\nBudget: Rs {proj['budget']:,}\nSpent: Rs {proj['spent']:,}")
    add_heading(doc, "Phase Summary", 2)
    for ph in proj["phases"]:
        add_para(doc, f"Phase {ph['phase']}: {ph['name']} — {ph['status']} (Due: {ph['end']})")
    add_heading(doc, "Key Risks", 2)
    add_para(doc,
        "1. Timeline risk if vendor tooling is delayed beyond Q2 2025.\n"
        "2. Budget overrun if additional testing cycles are required.\n"
        "3. Resource availability during peak production season.")
    add_heading(doc, "Next Steps", 2)
    pending = [ph for ph in proj["phases"] if ph["status"] == "Pending"]
    for ph in pending:
        add_para(doc, f"• Initiate {ph['name']} by {ph['end']}")
    fname = f"DOCX_01_Executive_Summary_{proj['id']}.docx"
    save_docx(doc, fname)


def docx_hr_intern_guide():
    doc = Document()
    doc.add_heading(CO["name"], 0)
    add_heading(doc, "Intern Onboarding and Project Guide — 2026")
    add_heading(doc, "About the Company", 2)
    add_para(doc,
        f"{CO['name']} was founded in {CO['founded']} and is India's leading pressure cooker manufacturer. "
        f"We operate {len(CO['factories'])} plants across Thane, Hoshiarpur, and Jaunpur, and export "
        f"to all six continents. Our mission: {CO['mission']}"
    )
    add_heading(doc, "Intern Projects — Summer 2026", 2)
    for i in INTERNS:
        add_para(doc, f"• {i['name']} ({i['college']}) — {i['project']} | Mentor: {i['mentor']}")
    add_heading(doc, "Leave Policy Summary", 2)
    add_para(doc,
        "Interns are entitled to 1 casual leave per month and up to 2 medical leaves per internship. "
        "Working hours are 9 AM to 6 PM, Monday to Friday. 90% attendance is mandatory for stipend. "
        "Full details in POL001 — Intern Leave and Conduct Policy."
    )
    add_heading(doc, "IT and Data Security", 2)
    add_para(doc,
        "All interns must adhere to POL004 — IT Security and Data Handling Policy. "
        "No company data may be stored on personal devices. All AI tools used for projects "
        "must process data locally. Internet access is restricted on project machines."
    )
    save_docx(doc, "DOCX_02_HR_Intern_Guide_2026.docx")


def docx_vendor_contract_summary(ven):
    doc = Document()
    doc.add_heading(CO["name"], 0)
    add_heading(doc, f"Vendor Contract Summary — {ven['name']}")
    add_heading(doc, "Vendor Information", 2)
    for k, v in ven.items():
        doc.add_paragraph(f"{k.replace('_',' ').title()}: {v}")
    add_heading(doc, "Contract Terms", 2)
    add_para(doc, f"Annual contract value: Rs {ven['contract_value']:,}")
    add_para(doc, "Payment terms: Net 30 days | Delivery: DDP to Hawkins plant")
    add_heading(doc, "Performance History", 2)
    add_para(doc, f"Performance score: {ven['rating']}/100 | Vendor since: {ven['since']}")
    add_heading(doc, "Approved By", 2)
    add_para(doc, f"{person('E003')['name']}, Procurement Head\n{person('E005')['name']}, Finance Director")
    fname = f"DOCX_03_Vendor_Contract_{ven['id']}.docx"
    save_docx(doc, fname)


def docx_finance_q4_budget():
    doc = Document()
    doc.add_heading(CO["name"], 0)
    add_heading(doc, "Finance — Q4 2025 Budget Review")
    add_heading(doc, "Executive Summary", 2)
    add_para(doc,
        "Q4 2025 ended with overall company performance within budget. "
        "Total revenue for Q4 was Rs 184 crores against a target of Rs 175 crores (+5.1%). "
        "Net profit for Q4: Rs 19.2 crores against target of Rs 17.8 crores.")
    add_heading(doc, "Department Budget Status", 2)
    depts = ["R&D", "Quality Control", "Procurement", "IT", "Operations", "Sales & Marketing"]
    for d in depts:
        budget = random.randint(400, 4500) * 10000
        spent  = int(budget * random.uniform(0.85, 0.98))
        add_para(doc, f"{d}: Budget Rs {budget:,} | Spent Rs {spent:,} | Variance Rs {budget-spent:,} (underspent)")
    add_heading(doc, "Project Budgets", 2)
    for p in PROJ:
        add_para(doc, f"{p['name']}: Approved Rs {p['budget']:,} | Spent Rs {p['spent']:,}")
    add_heading(doc, "Approved By", 2)
    fin = person("E005")
    add_para(doc, f"{fin['name']}, {fin['role']} | {fin['email']}")
    save_docx(doc, "DOCX_04_Finance_Q4_2025_Budget_Review.docx")


def docx_operations_plant_report():
    doc = Document()
    doc.add_heading(CO["name"], 0)
    add_heading(doc, "Operations — Thane Plant Monthly Report")
    add_heading(doc, "Production Summary", 2)
    for r in PROD_LOG[:6]:
        add_para(doc, f"{r['month']}: {r['units_produced']:,} units | Defect rate: {r['defect_rate_pct']}%")
    add_heading(doc, "Key Issues and Resolutions", 2)
    add_para(doc,
        "1. Line 2 coating oven recalibration completed in March 2025 — defect rate reduced.\n"
        "2. New gasket batch from SafeSeal Gaskets received and cleared QC.\n"
        "3. Capacity expansion for Contura line on track for Q3 2025.")
    add_heading(doc, "Staffing", 2)
    add_para(doc, f"Plant Manager: {person('E006')['name']}\nTotal plant staff: {DEPT['D006']['headcount']}")
    save_docx(doc, "DOCX_05_Operations_Thane_Plant_Monthly_Report.docx")


def docx_it_policy():
    doc = Document()
    doc.add_heading(CO["name"], 0)
    pol = next(p for p in POL if p["id"] == "POL004")
    add_heading(doc, pol["title"])
    add_para(doc, f"Effective: {pol['effective']} | Version: {pol['revision']}")
    add_heading(doc, "Policy Statement", 2)
    add_para(doc, pol["desc"])
    add_heading(doc, "AI Tools Guidelines", 2)
    add_para(doc,
        "All AI tools used for company projects must run locally on approved hardware. "
        "No company data may be sent to external AI APIs (e.g., OpenAI, Anthropic cloud). "
        "The approved stack for AI projects is: Ollama + LLaMA 3.1:8b + ChromaDB + BGE-M3. "
        "Exceptions require written approval from the IT Head.")
    add_heading(doc, "Data Classification", 2)
    add_para(doc,
        "Level 1 — Public: Product catalogs, price lists, cookbooks.\n"
        "Level 2 — Internal: Project reports, department budgets, HR policies.\n"
        "Level 3 — Confidential: Vendor contracts, audit findings, financial data.\n"
        "Level 4 — Restricted: R&D specifications, patent applications, board minutes.")
    add_heading(doc, "IT Head", 2)
    it = person("E008")
    add_para(doc, f"{it['name']} | {it['email']} | {it['phone']}")
    save_docx(doc, "DOCX_06_IT_Security_Data_Policy.docx")


# ─────────────────────────────────────────────────────────────────────────────
# EXCEL GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def save_xlsx(wb_xl, filename):
    path = os.path.join(RAW, "excel", filename)
    wb_xl.save(path)
    print(f"  [XLSX] {filename}")


def xlsx_production_log():
    wb_xl = openpyxl.Workbook()
    ws = wb_xl.active
    ws.title = "Production Log 2025"
    headers = ["Month", "Plant", "Product Line", "Units Produced", "Units Rejected", "Defect Rate %"]
    ws.append(headers)
    for r in PROD_LOG:
        ws.append([r["month"], r["plant"], r["product_line"],
                   r["units_produced"], r["units_rejected"], r["defect_rate_pct"]])
    ws2 = wb_xl.create_sheet("Plant Summary")
    ws2.append(["Plant", "Total Produced", "Total Rejected", "Avg Defect Rate %"])
    plants = {}
    for r in PROD_LOG:
        p = r["plant"]
        plants.setdefault(p, {"produced": 0, "rejected": 0, "count": 0})
        plants[p]["produced"]  += r["units_produced"]
        plants[p]["rejected"]  += r["units_rejected"]
        plants[p]["count"]     += 1
    for p, d in plants.items():
        ws2.append([p, d["produced"], d["rejected"],
                    round(d["rejected"]/d["produced"]*100, 2)])
    save_xlsx(wb_xl, "XLSX_01_Monthly_Production_Log_2025.xlsx")


def xlsx_vendor_invoice_tracker():
    wb_xl = openpyxl.Workbook()
    ws = wb_xl.active
    ws.title = "Invoice Tracker"
    ws.append(["Invoice No", "Vendor ID", "Vendor Name", "Amount (Rs)", "Date", "Status", "Approved By"])
    for inv in INV:
        v = vendor(inv["vendor_id"])
        ws.append([inv["invoice_no"], inv["vendor_id"], v["name"],
                   inv["amount"], inv["date"], inv["status"],
                   inv["approved_by"] or "Pending"])
    ws2 = wb_xl.create_sheet("Vendor Summary")
    ws2.append(["Vendor", "Total Invoiced", "Paid", "Pending"])
    vendor_totals = {}
    for inv in INV:
        v = vendor(inv["vendor_id"])
        vendor_totals.setdefault(v["name"], {"total": 0, "paid": 0, "pending": 0})
        vendor_totals[v["name"]]["total"] += inv["amount"]
        if inv["status"] == "Paid":
            vendor_totals[v["name"]]["paid"] += inv["amount"]
        else:
            vendor_totals[v["name"]]["pending"] += inv["amount"]
    for vn, d in vendor_totals.items():
        ws2.append([vn, d["total"], d["paid"], d["pending"]])
    save_xlsx(wb_xl, "XLSX_02_Vendor_Invoice_Tracker_2025.xlsx")


def xlsx_qc_defect_analysis():
    wb_xl = openpyxl.Workbook()
    ws = wb_xl.active
    ws.title = "QC Defect Analysis"
    ws.append(["Product Line", "Jan", "Feb", "Mar", "Apr", "May", "Jun",
               "Jul", "Aug", "Sep", "Oct", "Nov", "Dec", "Total", "Avg Monthly"])
    product_lines = list(set(r["product_line"] for r in PROD_LOG))
    for pl in product_lines:
        monthly = []
        for m in range(1, 13):
            relevant = [r for r in PROD_LOG if r["product_line"] == pl]
            if relevant:
                monthly.append(relevant[0]["units_rejected"])
            else:
                monthly.append(random.randint(50, 400))
        total = sum(monthly)
        avg   = round(total / 12, 1)
        ws.append([pl] + monthly + [total, avg])
    save_xlsx(wb_xl, "XLSX_03_QC_Defect_Analysis_2025.xlsx")


def xlsx_employee_roster():
    wb_xl = openpyxl.Workbook()
    ws = wb_xl.active
    ws.title = "Employee Roster"
    ws.append(["Employee ID", "Name", "Role", "Department", "Email", "Phone", "Location", "Joined"])
    for e in PPL:
        dept_name = DEPT[e["dept"]]["name"]
        ws.append([e["id"], e["name"], e["role"], dept_name,
                   e["email"], e["phone"], e["location"], e["joined"]])
    # Add 20 Faker-generated employees
    for i in range(20):
        eid = f"E{100+i:03d}"
        dept = random.choice(list(DEPT.values()))
        ws.append([eid, fake.name(), fake.job()[:30], dept["name"],
                   fake.email(), fake.phone_number()[:15],
                   random.choice(["Mumbai HQ","Thane Plant","Hoshiarpur","Jaunpur"]),
                   fake.date_between(start_date="-8y", end_date="-1y").isoformat()])
    ws2 = wb_xl.create_sheet("Interns 2026")
    ws2.append(["ID", "Name", "College", "Department", "Mentor", "Project", "Start", "End"])
    for i in INTERNS:
        ws2.append([i["id"], i["name"], i["college"],
                    DEPT[i["dept"]]["name"], i["mentor"],
                    i["project"], i["start"], i["end"]])
    save_xlsx(wb_xl, "XLSX_04_Employee_Department_Roster.xlsx")


def xlsx_project_budget_tracker():
    wb_xl = openpyxl.Workbook()
    ws = wb_xl.active
    ws.title = "Project Budget Tracker"
    ws.append(["Project ID", "Project Name", "Department", "Lead",
               "Budget (Rs)", "Spent (Rs)", "Remaining (Rs)", "Utilisation %", "Status"])
    for p in PROJ:
        remaining    = p["budget"] - p["spent"]
        utilisation  = round(p["spent"] / p["budget"] * 100, 1)
        ws.append([p["id"], p["name"], DEPT[p["dept"]]["name"], p["lead"],
                   p["budget"], p["spent"], remaining, utilisation, p["status"]])
    ws2 = wb_xl.create_sheet("Phase Details")
    ws2.append(["Project", "Phase No", "Phase Name", "Status", "Target End"])
    for p in PROJ:
        for ph in p["phases"]:
            ws2.append([p["name"], ph["phase"], ph["name"], ph["status"], ph["end"]])
    save_xlsx(wb_xl, "XLSX_05_Project_Budget_Tracker.xlsx")


def xlsx_price_list():
    wb_xl = openpyxl.Workbook()
    ws = wb_xl.active
    ws.title = "Pressure Cookers"
    ws.append(["ID", "Brand", "Model", "Capacity (L)", "Product Code", "MRP (Rs)", "Induction", "Description"])
    for p in PC:
        ws.append([p["id"], p["brand"], p["model"], p.get("capacity_L",""),
                   p["code"], p["mrp"], "Yes" if p["induction"] else "No", p["desc"]])
    ws2 = wb_xl.create_sheet("Cookware")
    ws2.append(["ID", "Brand", "Type", "Product Code", "MRP (Rs)", "Induction", "Description"])
    for c in CW:
        ws2.append([c["id"], c["brand"], c["type"], c["code"], c["mrp"],
                    "Yes" if c["induction"] else "No", c["desc"]])
    ws3 = wb_xl.create_sheet("Kitchen Electricals")
    ws3.append(["ID", "Brand", "Type", "Code", "MRP (Rs)", "Description"])
    for k in KE:
        ws3.append([k["id"], k["brand"], k["type"], k["code"], k["mrp"], k["desc"]])
    save_xlsx(wb_xl, "XLSX_06_Full_Price_List_June_2026.xlsx")


# ─────────────────────────────────────────────────────────────────────────────
# EMAIL GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def make_email(from_p, to_p, subject, body, date=None):
    msg = MIMEMultipart()
    msg["From"]    = f"{from_p['name']} <{from_p['email']}>"
    msg["To"]      = f"{to_p['name']} <{to_p['email']}>"
    msg["Subject"] = subject
    msg["Date"]    = date or rand_date()
    msg.attach(MIMEText(body, "plain"))
    return msg


def save_email(msg, filename):
    path = os.path.join(RAW, "emails", filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(msg.as_string())
    print(f"  [EML]  {filename}")


def generate_emails():
    emails = [
        # Project Aurora emails
        (person("E001"), person("E002"),
         "Project Aurora — Phase 2 approved, Phase 3 starting",
         f"Hi {person('E002')['name'].split()[0]},\n\nI am pleased to inform you that the Phase 2 prototype "
         f"of Project Aurora has passed the 15-bar pressure test (Audit AUD002, Jan 20 2025). "
         f"The smart valve design is performing as per spec.\n\n"
         f"We are now moving into Phase 3 — Pressure Testing and BIS certification. "
         f"I will need QC support for the batch testing at Thane plant. "
         f"Can we schedule a meeting this week?\n\nRegards,\n{person('E001')['name']}",
         "2025-01-22"),

        (person("E001"), person("E005"),
         "Project Aurora — Budget utilisation update Q1 2025",
         f"Dear {person('E005')['name'].split()[0]},\n\nProject Aurora budget update for Q1 2025:\n"
         f"Approved budget: Rs 45,00,000\nSpent to date: Rs 28,70,000\nRemaining: Rs 16,30,000\n\n"
         f"We are on track. The main upcoming cost is BIS certification fees and "
         f"external testing lab charges. Estimated Rs 4,50,000 in Q2.\n\n"
         f"Please let me know if you need a detailed breakdown.\n\nRegards,\n{person('E001')['name']}",
         "2025-03-31"),

        # Project Falcon emails
        (person("E002"), person("E001"),
         "Project Falcon — Final audit report uploaded",
         f"Hi {person('E001')['name'].split()[0]},\n\nProject Falcon QC audit is now complete. "
         f"The final report has been uploaded to the shared drive.\n\n"
         f"Key findings:\n- 3 critical defects found in AC-3L batch QC-2024-041 (gasket misfit on 4.2% units)\n"
         f"- Jaunpur line 3 coating thickness variance resolved (Audit AUD004)\n"
         f"- All remediation actions completed by December 2024\n\n"
         f"Overall the AlphaCook line is now compliant with POL003 QC standards.\n\n"
         f"Best,\n{person('E002')['name']}",
         "2025-01-05"),

        (person("E002"), person("E006"),
         "QC Failure Alert — Batch HC-2025-012 Contura",
         f"Dear {person('E006')['name'].split()[0]},\n\nUrgent: QC inspection has flagged batch "
         f"HC-2025-012 (Hawkins Contura 3L) at Thane plant.\n\n"
         f"Issue: Ceramic nonstick coating adhesion below specification on 6 units (2% defect rate).\n"
         f"Root cause: Temperature variance on coating line 2 — sensors drifted by 8 degrees.\n"
         f"Action required: Halt production on line 2 until recalibration is complete.\n\n"
         f"Please confirm batch hold and arrange sensor recalibration today.\n\n"
         f"Regards,\n{person('E002')['name']}",
         "2025-02-18"),

        # Vendor / Procurement emails
        (person("E003"), person("E004"),
         "Vendor Presstek — Contract renewal approval needed",
         f"Hi {person('E004')['name'].split()[0]},\n\nPresstek Components Ltd contract is due for renewal "
         f"(current contract expires Dec 2025).\n\n"
         f"Vendor performance score: 91/100 (Audit AUD003)\n"
         f"Recommended renewal: YES\n"
         f"Proposed contract value: Rs 18,00,000 for 2026\n\n"
         f"Please process the approval so we can issue the PO by end of month.\n\n"
         f"Thanks,\n{person('E003')['name']}",
         "2025-02-12"),

        (person("E005"), person("E003"),
         "Invoice INV-2025-0034 from SteelForm approved",
         f"Dear {person('E003')['name'].split()[0]},\n\nInvoice INV-2025-0034 from SteelForm Industries "
         f"for Rs 9,10,000 has been reviewed and approved for payment.\n\n"
         f"Payment will be processed within 5 working days as per our Net-30 terms.\n\n"
         f"Please ensure the delivery challan and QC clearance certificate are filed.\n\n"
         f"Regards,\n{person('E005')['name']}",
         "2025-02-05"),

        (person("E003"), person("E005"),
         "New vendor GermanCoat GmbH — Onboarding complete",
         f"Hi {person('E005')['name'].split()[0]},\n\nGood news — GermanCoat GmbH (India) has completed "
         f"the vendor onboarding process as per POL002.\n\n"
         f"Site visit: Completed Jan 2025\nTrial order: Cleared QC (score 97/100)\n"
         f"Contract value: Rs 38,00,000 per annum\n"
         f"They will supply high-quality German ceramic coating for the Ceramic Nonstick range.\n\n"
         f"Please arrange first payment against INV-2025-0037.\n\n"
         f"Regards,\n{person('E003')['name']}",
         "2025-03-11"),

        # HR / Intern emails
        (person("E004"), None,
         "Intern Leave Policy — Updated effective January 2025",
         f"Dear All Interns,\n\nPlease note that the Intern Leave and Conduct Policy (POL001) "
         f"has been updated effective January 1, 2025 (version 3.0).\n\n"
         f"Key changes:\n"
         f"1. Medical leave increased from 1 day to 2 days per internship period\n"
         f"2. Work from home now permitted with mentor approval\n"
         f"3. Completion bonus increased to Rs 5,000\n\n"
         f"Please read the full policy document available on the HR portal.\n\n"
         f"For queries contact: {person('E004')['email']}\n\n"
         f"Best regards,\n{person('E004')['name']}\nHR Manager",
         "2025-01-02"),

        (person("E004"), person("E008"),
         "Intern Sambodh Gupta — Project assignment confirmation",
         f"Dear {person('E008')['name'].split()[0]},\n\nThis is to confirm that intern "
         f"Sambodh Gupta from IIIT Manipur has been assigned to your department for the "
         f"Summer 2026 internship.\n\n"
         f"Project: Local Enterprise Knowledge Assistant (Project Horizon, Phase 2)\n"
         f"Duration: June 1 to July 31, 2026\n"
         f"Stipend: Rs 15,000/month\n\n"
         f"Please ensure onboarding checklist is completed on Day 1.\n\n"
         f"Regards,\n{person('E004')['name']}",
         "2026-05-28"),

        # IT emails
        (person("E008"), person("E001"),
         "Project Horizon — ERP assessment findings",
         f"Dear {person('E001')['name'].split()[0]},\n\nThe ERP assessment (Audit AUD006) for "
         f"Project Horizon is complete.\n\n"
         f"Recommendation: SAP S/4HANA upgrade over Oracle ERP\n"
         f"Estimated cost: Rs 85,00,000\n"
         f"Implementation timeline: 18 months\n\n"
         f"The AI Knowledge Assistant sub-project is progressing well — "
         f"intern Sambodh Gupta is building the prototype using LLaMA 3.1:8b with "
         f"ChromaDB for local vector search. No data leaves the network.\n\n"
         f"Will present to leadership next month.\n\nRegards,\n{person('E008')['name']}",
         "2025-07-15"),

        (person("E008"), person("E005"),
         "IT Budget Q2 2025 — Project Horizon spend update",
         f"Dear {person('E005')['name'].split()[0]},\n\nProject Horizon IT budget update for Q2 2025:\n\n"
         f"ERP assessment: Rs 2,20,000 (completed)\n"
         f"AI prototype hardware/software: Rs 80,000\n"
         f"IoT sensor pilot (3 machines): Rs 1,20,000\n"
         f"Total spent Q2: Rs 4,20,000\n"
         f"Budget remaining for FY: Rs 13,80,000\n\n"
         f"On track. No budget revision needed at this stage.\n\n"
         f"Regards,\n{person('E008')['name']}",
         "2025-07-31"),

        # Export emails
        (person("E010"), person("E007"),
         "Export Q1 2025 — USA and Europe performance",
         f"Hi {person('E007')['name'].split()[0]},\n\nExport performance for Q1 2025:\n\n"
         f"USA: 12,400 units shipped (Futura HA, UL-listed Classic)\n"
         f"Europe: 8,200 units shipped (Futura SS, CE certified range)\n"
         f"Middle East: 6,100 units shipped (Classic, Bigboy)\n"
         f"Total export revenue Q1: Rs 22.4 crores\n\n"
         f"All shipments cleared customs with valid UL/CE/BIS documentation.\n\n"
         f"Regards,\n{person('E010')['name']}",
         "2025-04-10"),

        # Operations emails
        (person("E006"), person("E002"),
         "Thane Plant — Line 2 recalibration completed",
         f"Dear {person('E002')['name'].split()[0]},\n\nLine 2 coating oven recalibration has been "
         f"completed as requested following the Contura batch HC-2025-012 QC alert.\n\n"
         f"Temperature variance: Reduced from 8 degrees to within 1.5 degrees of spec.\n"
         f"Test run completed: 50 units — 0 defects found.\n"
         f"Line 2 is cleared for full production.\n\n"
         f"Batch hold on HC-2025-012 (6 affected units) has been quarantined for rework.\n\n"
         f"Regards,\n{person('E006')['name']}",
         "2025-02-22"),

        (person("E006"), person("E003"),
         "Raw material stock alert — Aluminium sheets running low",
         f"Dear {person('E003')['name'].split()[0]},\n\nUrgent stock alert from Thane plant.\n\n"
         f"Current aluminium sheet stock: 8 days of production remaining.\n"
         f"Vendor: SteelForm Industries (V002)\n"
         f"Last order: INV-2025-0034 (delivered Feb 3)\n"
         f"Required: Minimum 30-day buffer stock\n\n"
         f"Please raise a purchase order immediately to avoid production stoppage.\n\n"
         f"Regards,\n{person('E006')['name']}",
         "2025-03-05"),

        # Customer service emails
        (person("E009"), person("E002"),
         "Customer complaint — Pressure cooker lid safety issue",
         f"Dear {person('E002')['name'].split()[0]},\n\nWe have received 3 customer complaints "
         f"in the past week regarding the Hawkins Classic 4L lid not locking properly.\n\n"
         f"All 3 complaints are from the same retailer (Big Bazaar, Pune).\n"
         f"Product code: CL 40 | Batch: CL-2025-055\n\n"
         f"Please arrange a QC investigation of this batch immediately.\n"
         f"I have put a temporary hold on returns until we hear from QC.\n\n"
         f"Regards,\n{person('E009')['name']}",
         "2025-04-18"),

        # Cookware product email
        (person("E001"), person("E007"),
         "New product launch — Hawkins Cerenity SS range",
         f"Hi {person('E007')['name'].split()[0]},\n\nProject Cerenity is progressing well. "
         f"The prototype spring-finish adhesion test has passed (Audit AUD005, March 2025).\n\n"
         f"Planned launch: Q1 2026\n"
         f"Target segment: Premium SS pressure cooker market\n"
         f"Key USPs: Spring collection finish, induction compatible, 5-year guarantee\n"
         f"Suggested MRP: Rs 3,500 (3L model)\n\n"
         f"Please start planning the marketing campaign. "
         f"I'll share the product spec sheet by end of week.\n\n"
         f"Regards,\n{person('E001')['name']}",
         "2025-03-15"),

        # Additional emails to reach 20
        (person("E003"), person("E001"),
         "Vendor SafeSeal — Gasket replacement batch cleared",
         f"Hi {person('E001')['name'].split()[0]},\n\nGood news — the replacement gasket batch "
         f"from SafeSeal Gaskets (V003) has been received and cleared QC inspection.\n\n"
         f"The dimensional variance issue that caused the Futura SS 5L batch rejection "
         f"has been resolved. SafeSeal has updated their tooling.\n\n"
         f"Production on Futura SS line can resume at full capacity.\n\n"
         f"Regards,\n{person('E003')['name']}",
         "2025-02-28"),

        (person("E007"), person("E010"),
         "Export pricelist update — Effective June 2026",
         f"Dear {person('E010')['name'].split()[0]},\n\nPlease note that the Hawkins pricelist "
         f"has been updated effective June 1, 2026.\n\n"
         f"Key changes for export markets:\n"
         f"- Hawkins Classic 2L (CL 20): Rs 1,425 (unchanged)\n"
         f"- Futura HA 4L: Rs 3,999 (5% increase)\n"
         f"- Cerenity SS 3L (new): Rs 3,500\n"
         f"- Electronic Kettle FKTA 1: Rs 2,750 (unchanged)\n\n"
         f"Please update your export quotes accordingly.\n\n"
         f"Regards,\n{person('E007')['name']}",
         "2026-06-01"),

        (person("E004"), person("E003"),
         "Vendor site visit — Schedule confirmation for Tristar Metals",
         f"Dear {person('E003')['name'].split()[0]},\n\nAs per the vendor onboarding process (POL002), "
         f"the site visit to Tristar Metals Pvt Ltd (V004) in Ahmedabad has been scheduled for "
         f"March 20, 2025.\n\n"
         f"Attendees: {person('E003')['name']} (Procurement), {person('E002')['name']} (QC)\n"
         f"Travel has been approved under the Travel and Expense Reimbursement Policy (POL005).\n"
         f"Hotel: Marriott Ahmedabad (booked by HR)\n\n"
         f"Please carry the vendor evaluation checklist.\n\n"
         f"Regards,\n{person('E004')['name']}",
         "2025-03-10"),

        (person("E005"), person("E008"),
         "Project Horizon — Budget release for Phase 2",
         f"Dear {person('E008')['name'].split()[0]},\n\nI am pleased to confirm that the budget "
         f"for Project Horizon Phase 2 (AI Knowledge Assistant) has been released.\n\n"
         f"Amount: Rs 6,00,000\n"
         f"Purpose: Intern stipend, hardware, software licences, testing\n"
         f"Valid through: July 31, 2026\n\n"
         f"Please ensure all expenses are submitted with receipts within 7 days of incurring them.\n\n"
         f"Regards,\n{person('E005')['name']}",
         "2026-06-02"),
    ]

    # Handle the broadcast email (to=None)
    for idx, e in enumerate(emails):
        if e[1] is None:
            to_p = {"name": "All Interns", "email": "interns@hawkins.com"}
        else:
            to_p = e[1]
        msg = make_email(e[0], to_p, e[2], e[3], e[4])
        save_email(msg, f"EMAIL_{idx+1:02d}_{e[2][:40].replace(' ','_').replace('—','')}.eml")


# ─────────────────────────────────────────────────────────────────────────────
# SQL DATABASE GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def generate_sql():
    db_path = os.path.join(RAW, "sql", "hawkins.db")
    conn = sqlite3.connect(db_path)
    c = conn.cursor()

    c.execute("DROP TABLE IF EXISTS audit_records")
    c.execute("""
        CREATE TABLE audit_records (
            audit_id        TEXT PRIMARY KEY,
            project_id      TEXT,
            plant           TEXT,
            date            TEXT,
            auditor         TEXT,
            finding         TEXT,
            status          TEXT,
            resolution_date TEXT
        )
    """)
    for a in AUD:
        c.execute("INSERT INTO audit_records VALUES (?,?,?,?,?,?,?,?)",
                  (a["audit_id"], a["project_id"], a["plant"], a["date"],
                   a["auditor"], a["finding"], a["status"], a.get("resolution_date")))

    c.execute("DROP TABLE IF EXISTS vendor_invoices")
    c.execute("""
        CREATE TABLE vendor_invoices (
            invoice_no  TEXT PRIMARY KEY,
            vendor_id   TEXT,
            vendor_name TEXT,
            amount      INTEGER,
            date        TEXT,
            status      TEXT,
            approved_by TEXT
        )
    """)
    for inv in INV:
        v = vendor(inv["vendor_id"])
        c.execute("INSERT INTO vendor_invoices VALUES (?,?,?,?,?,?,?)",
                  (inv["invoice_no"], inv["vendor_id"], v["name"],
                   inv["amount"], inv["date"], inv["status"],
                   inv.get("approved_by") or "Pending"))

    c.execute("DROP TABLE IF EXISTS production_log")
    c.execute("""
        CREATE TABLE production_log (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            month           TEXT,
            plant           TEXT,
            product_line    TEXT,
            units_produced  INTEGER,
            units_rejected  INTEGER,
            defect_rate_pct REAL
        )
    """)
    for r in PROD_LOG:
        c.execute("INSERT INTO production_log (month,plant,product_line,units_produced,units_rejected,defect_rate_pct) VALUES (?,?,?,?,?,?)",
                  (r["month"], r["plant"], r["product_line"],
                   r["units_produced"], r["units_rejected"], r["defect_rate_pct"]))

    c.execute("DROP TABLE IF EXISTS employees")
    c.execute("""
        CREATE TABLE employees (
            emp_id    TEXT PRIMARY KEY,
            name      TEXT,
            role      TEXT,
            dept      TEXT,
            email     TEXT,
            location  TEXT,
            joined    TEXT
        )
    """)
    for e in PPL:
        c.execute("INSERT INTO employees VALUES (?,?,?,?,?,?,?)",
                  (e["id"], e["name"], e["role"],
                   DEPT[e["dept"]]["name"], e["email"], e["location"], e["joined"]))

    c.execute("DROP TABLE IF EXISTS projects")
    c.execute("""
        CREATE TABLE projects (
            project_id  TEXT PRIMARY KEY,
            name        TEXT,
            dept        TEXT,
            lead        TEXT,
            status      TEXT,
            budget      INTEGER,
            spent       INTEGER,
            start_date  TEXT,
            end_date    TEXT
        )
    """)
    for p in PROJ:
        c.execute("INSERT INTO projects VALUES (?,?,?,?,?,?,?,?,?)",
                  (p["id"], p["name"], DEPT[p["dept"]]["name"], p["lead"],
                   p["status"], p["budget"], p["spent"], p["start"], p["end"]))

    conn.commit()
    conn.close()
    print(f"  [SQL]  hawkins.db — 5 tables created")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main():
    print("\n" + "="*60)
    print("  HAWKINS DATA GENERATOR")
    print("="*60)

    # ── PDFs ──────────────────────────────────────────────────────
    print("\n[1] Generating PDFs...")
    pdf_company_overview()
    for p in PROJ:
        pdf_project_report(p)
    for v in VEN:
        pdf_vendor_evaluation(v)
    pdf_qc_inspection_report()
    for pol in POL:
        pdf_hr_policy(pol)
    pdf_annual_production_review()
    pdf_export_report()
    for p in PC:
        pdf_product_spec(p, "PC")
    for c in CW:
        pdf_product_spec(c, "CW")
    pdf_it_project_report()
    pdf_service_centre_report()
    pdf_finance_budget()
    pdf_intern_onboarding()
    pdf_safety_compliance()
    pdf_product_catalog_summary()
    for v in VEN:
        pdf_vendor_contract(v)
    pdf_audit_summary()

    # ── Word docs ─────────────────────────────────────────────────
    print("\n[2] Generating Word documents...")
    for p in PROJ:
        docx_executive_summary(p)
    docx_hr_intern_guide()
    for v in VEN:
        docx_vendor_contract_summary(v)
    docx_finance_q4_budget()
    docx_operations_plant_report()
    docx_it_policy()

    # ── Excel ─────────────────────────────────────────────────────
    print("\n[3] Generating Excel files...")
    xlsx_production_log()
    xlsx_vendor_invoice_tracker()
    xlsx_qc_defect_analysis()
    xlsx_employee_roster()
    xlsx_project_budget_tracker()
    xlsx_price_list()

    # ── Emails ────────────────────────────────────────────────────
    print("\n[4] Generating emails...")
    generate_emails()

    # ── SQL ───────────────────────────────────────────────────────
    print("\n[5] Generating SQL database...")
    generate_sql()

    # ── Summary ───────────────────────────────────────────────────
    pdf_count   = len(os.listdir(os.path.join(RAW, "pdfs")))
    docx_count  = len(os.listdir(os.path.join(RAW, "docx")))
    excel_count = len(os.listdir(os.path.join(RAW, "excel")))
    email_count = len(os.listdir(os.path.join(RAW, "emails")))
    sql_count   = len(os.listdir(os.path.join(RAW, "sql")))

    print("\n" + "="*60)
    print("  GENERATION COMPLETE")
    print("="*60)
    print(f"  PDFs:   {pdf_count}")
    print(f"  DOCX:   {docx_count}")
    print(f"  Excel:  {excel_count}")
    print(f"  Emails: {email_count}")
    print(f"  SQL:    {sql_count}")
    print(f"  TOTAL:  {pdf_count+docx_count+excel_count+email_count+sql_count} files")
    print("="*60)


if __name__ == "__main__":
    main()